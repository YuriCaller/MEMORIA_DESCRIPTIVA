# -*- coding: utf-8 -*-
"""
Generación del documento Word con formato profesional justificado.
VERSIÓN 2.1 - Soporte para modo atlas (un documento por feature/polígono).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


# ===========================================================================
# HELPERS DE FORMATO
# ===========================================================================

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def _set_cell_borders(cell, color="AAAAAA", size=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement('w:{}'.format(side))
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(size))
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _add_bottom_border(p, color='17375E', size='6'):
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),size)
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

def _para_just(doc, text, size=11, bold=False, sp_before=0, sp_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(sp_before)
    p.paragraph_format.space_after  = Pt(sp_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text); r.font.size = Pt(size)
    r.font.bold = bold; r.font.name = 'Arial'
    return p

def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Arial'; r.font.bold = True
    if level == 1:
        r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x17,0x37,0x5E)
        _add_bottom_border(p, '17375E', '6')
    elif level == 2:
        r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x2E,0x6E,0x3E)
        _add_bottom_border(p, '2E6E3E', '4')
    else:
        r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x40,0x40,0x40)
    return p

def _fmt_hdr(row, bg='17375E', fg='FFFFFF'):
    for cell in row.cells:
        _set_cell_bg(cell, bg); _set_cell_borders(cell, bg, 4)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            for r in para.runs:
                r.font.bold=True; r.font.name='Arial'; r.font.size=Pt(10)
                r.font.color.rgb = RGBColor(int(fg[0:2],16),int(fg[2:4],16),int(fg[4:6],16))

def _fmt_row(row, alt=False):
    bg = 'EAF4EA' if alt else 'FFFFFF'
    for cell in row.cells:
        _set_cell_bg(cell, bg); _set_cell_borders(cell, 'CCCCCC', 2)
        for para in cell.paragraphs:
            para.paragraph_format.space_before=Pt(2); para.paragraph_format.space_after=Pt(2)
            for r in para.runs: r.font.name='Arial'; r.font.size=Pt(9)

def _fnum(v, fmt=',.4f'):
    try: return format(float(v), fmt)
    except: return str(v)

_MESES = {'January':'enero','February':'febrero','March':'marzo','April':'abril',
          'May':'mayo','June':'junio','July':'julio','August':'agosto',
          'September':'setiembre','October':'octubre','November':'noviembre','December':'diciembre'}

def _fecha_es():
    s = datetime.now().strftime('%d de %B de %Y')
    for en,es in _MESES.items(): s = s.replace(en,es)
    return s


# ===========================================================================
# FUNCIÓN PRINCIPAL
# ===========================================================================

def generar_documento_word(datos_formulario, datos_procesados, sufijo_archivo=None):
    """
    Genera un documento Word de Memoria Descriptiva.

    Args:
        datos_formulario : dict con todos los datos del formulario
        datos_procesados : dict con vértices, área, perímetro, colindantes, etc.
        sufijo_archivo   : string para diferenciar archivos en modo atlas (ej. nombre del predio)

    Returns:
        Ruta del archivo generado
    """
    doc = Document()

    for section in doc.sections:
        section.page_width   = Cm(21.0); section.page_height  = Cm(29.7)
        section.top_margin   = Cm(2.5);  section.bottom_margin = Cm(2.5)
        section.left_margin  = Cm(3.0);  section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    _encabezado(doc, datos_formulario)
    _seccion_solicitante(doc, datos_formulario.get('solicitante', {}))
    _seccion_generalidades(doc, datos_formulario.get('generalidades', ''))
    _seccion_ubicacion(doc, datos_formulario.get('ubicacion', {}))
    _seccion_colindantes(doc, datos_procesados.get('colindantes', {}))
    _seccion_tecnica(doc, datos_procesados)
    _seccion_mapa(doc, datos_formulario.get('info_mapa', {}))
    _seccion_firma(doc, datos_formulario.get('solicitante', {}))

    # Determinar ruta de salida
    output_base = datos_formulario['output_file']
    if sufijo_archivo:
        base, ext = os.path.splitext(output_base)
        # Limpiar caracteres no válidos para nombre de archivo
        sufijo_limpio = "".join(c for c in sufijo_archivo if c.isalnum() or c in (' ','_','-')).strip()
        sufijo_limpio = sufijo_limpio[:40]  # máximo 40 chars
        output_path = "{}_{}{}".format(base, sufijo_limpio, ext)
    else:
        output_path = output_base

    doc.save(output_path)
    return output_path


# ===========================================================================
# SECCIONES
# ===========================================================================

def _encabezado(doc, datos):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4)
    r = p.add_run('MEMORIA DESCRIPTIVA')
    r.font.name='Arial'; r.font.size=Pt(16); r.font.bold=True
    r.font.color.rgb = RGBColor(0x17,0x37,0x5E)

    sector = datos.get('ubicacion',{}).get('sector','')
    nombre_predio = datos.get('nombre_predio','')
    subtitulo = nombre_predio or sector
    if subtitulo:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(2)
        r2 = p2.add_run('DEL PREDIO: {}'.format(subtitulo.upper()))
        r2.font.name='Arial'; r2.font.size=Pt(12); r2.font.bold=True
        r2.font.color.rgb = RGBColor(0x2E,0x6E,0x3E)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before=Pt(4); p3.paragraph_format.space_after=Pt(10)
    _add_bottom_border(p3,'17375E','12')


def _seccion_solicitante(doc, dat):
    _heading(doc, 'I.   DATOS DEL SOLICITANTE', 1)
    for label, val in [
        ('Nombre y Apellidos : ', (dat.get('nombre','') or 'No especificado').upper()),
        ('D.N.I.                      : ', dat.get('dni','') or 'No especificado'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        r1 = p.add_run(label); r1.font.name='Arial'; r1.font.size=Pt(11)
        r2 = p.add_run(val);   r2.font.name='Arial'; r2.font.size=Pt(11); r2.font.bold=True


def _seccion_generalidades(doc, texto):
    _heading(doc, 'II.  GENERALIDADES', 1)
    if not texto or not texto.strip():
        texto = ("La presente Memoria Descriptiva tiene por finalidad describir las "
                 "características técnicas del predio materia del presente trámite, "
                 "determinando sus linderos, medidas perimétricas, área total, colindantes "
                 "y demás aspectos técnicos que permitan su correcta identificación y "
                 "ubicación en el territorio nacional, de acuerdo con las normas "
                 "técnicas vigentes del SERFOR.")
    for i, par in enumerate([p.strip() for p in texto.split('\n') if p.strip()]):
        _para_just(doc, par, sp_before=4 if i==0 else 2, sp_after=4)


def _seccion_ubicacion(doc, dat):
    _heading(doc, 'III. UBICACIÓN', 1)
    sector=dat.get('sector',''); distrito=dat.get('distrito','')
    provincia=dat.get('provincia',''); depto=dat.get('departamento','')
    zona=dat.get('zona','')
    partes = []
    if sector:    partes.append('el sector denominado {}'.format(sector))
    if distrito:  partes.append('el Distrito de {}'.format(distrito))
    if provincia: partes.append('la Provincia de {}'.format(provincia))
    if depto:     partes.append('el Departamento de {}'.format(depto))
    if partes:
        txt = ('El predio materia de la presente memoria descriptiva se encuentra '
               'ubicado en {}, República del Perú.'.format(', '.join(partes)))
    else:
        txt = 'El predio se encuentra ubicado en la República del Perú.'
    _para_just(doc, txt, sp_before=4, sp_after=6)

    items = [(k,v) for k,v in [
        ('Sector / Localidad', sector), ('Zona UTM', zona),
        ('Distrito', distrito), ('Provincia', provincia),
        ('Departamento', depto), ('País','Perú')] if v]
    if items:
        t = doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]; hdr.cells[0].text='DESCRIPCIÓN'; hdr.cells[1].text='DETALLE'
        _fmt_hdr(hdr)
        for i,(k,v) in enumerate(items):
            row=t.add_row(); row.cells[0].text=k; row.cells[1].text=str(v)
            _fmt_row(row, i%2==1)
            for c in row.cells:
                c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
                for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)


def _seccion_colindantes(doc, dat):
    _heading(doc, 'IV.  COLINDANTES', 1)
    orden = ['NORTE','SUR','ESTE','OESTE']
    partes = []
    for lado in orden:
        info = dat.get(lado, {})
        nombre = info.get('nombre','Terrenos del Estado') if isinstance(info,dict) else str(info)
        partes.append('por el {} con {}'.format(lado.capitalize(), nombre))
    if partes:
        txt = ('El predio colinda: ' + '; '.join(partes[:-1]) +
               (' y ' + partes[-1] if len(partes)>1 else partes[0]) + '.')
        _para_just(doc, txt, sp_before=4, sp_after=6)

    t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=t.rows[0]; hdr.cells[0].text='LADO'; hdr.cells[1].text='COLINDANTE'; hdr.cells[2].text='OBSERVACIÓN'
    _fmt_hdr(hdr)
    for i,lado in enumerate(orden):
        info=dat.get(lado,{})
        nombre = info.get('nombre','Terrenos del Estado') if isinstance(info,dict) else str(info)
        obs    = info.get('observacion','') if isinstance(info,dict) else ''
        row=t.add_row(); row.cells[0].text='POR EL {}'.format(lado)
        row.cells[1].text=str(nombre); row.cells[2].text=str(obs)
        _fmt_row(row, i%2==1)
        row.cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        try: row.cells[0].paragraphs[0].runs[0].font.bold=True
        except: pass
        for c in row.cells:
            for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)


def _seccion_tecnica(doc, dp):
    _heading(doc, 'V.   INFORMACIÓN TÉCNICA DEL PREDIO', 1)
    vertices    = dp.get('vertices', [])
    area        = dp.get('area', 0)
    perimetro   = dp.get('perimetro', 0)
    desc_lin    = dp.get('descripcion_linderos', '')
    fuente_area = dp.get('fuente_area', 'calculada')

    # 5.1 Linderos
    _heading(doc, '5.1.   Linderos y Medidas Perimétricas', 2)
    _para_just(doc,
        'El predio se enmarca con las siguientes medidas perimétricas, expresadas '
        'en el Sistema UTM (Universal Transversal de Mercator), con coordenadas en metros:',
        sp_before=4, sp_after=4)
    if desc_lin:
        _para_just(doc, desc_lin, sp_before=2, sp_after=6)

    # 5.2 Cuadro técnico
    _heading(doc, '5.2.   Cuadro Técnico de Vértices', 2)
    _para_just(doc,
        'A continuación, se presenta el cuadro técnico con las coordenadas de los '
        'vértices del predio, junto con las distancias y azimuts de cada lado:',
        sp_before=4, sp_after=6)

    if vertices:
        t=doc.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]
        for i,h in enumerate(['VÉRTICE','LADO','ESTE (m)','NORTE (m)','DISTANCIA (m)','AZIMUT (°)']):
            hdr.cells[i].text=h
        _fmt_hdr(hdr)
        for idx,v in enumerate(vertices):
            row=t.add_row()
            row.cells[0].text = str(v.get('vertice','V-{:02d}'.format(idx+1)))
            row.cells[1].text = str(v.get('lado',''))
            row.cells[2].text = _fnum(v.get('este',0),',.4f')
            row.cells[3].text = _fnum(v.get('norte',0),',.4f')
            row.cells[4].text = _fnum(v.get('distancia',0),'.2f')
            row.cells[5].text = _fnum(v.get('azimut',0),'.4f')
            _fmt_row(row, idx%2==1)
            for c in row.cells:
                c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(9)
    else:
        _para_just(doc,
            'No se encontraron vértices en la capa de puntos seleccionada. '
            'Verifique que la capa contiene los datos requeridos.',
            sp_before=4, sp_after=6)

    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # 5.3 Área y Perímetro
    _heading(doc, '5.3.   Área y Perímetro', 2)
    try:
        ah=float(area); am2=round(ah*10000,2)
        texto_area = ('El predio tiene una superficie total de {:,.4f} hectáreas '
                      '({:,.2f} m²) y un perímetro de {:,.2f} metros lineales. '
                      '[Fuente: {}]'.format(ah, am2, float(perimetro), fuente_area))
    except Exception:
        texto_area = 'Los datos de área y perímetro se calcularán con base en la geometría del polígono.'
    _para_just(doc, texto_area, sp_before=4, sp_after=6)

    t2=doc.add_table(rows=2,cols=2); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; t2.style='Table Grid'
    hdr2=t2.rows[0]; hdr2.cells[0].text='ÁREA TOTAL'; hdr2.cells[1].text='PERÍMETRO TOTAL'
    _fmt_hdr(hdr2, bg='2E6E3E')
    dr=t2.rows[1]
    dr.cells[0].text='{} ha'.format(_fnum(area,',.4f'))
    dr.cells[1].text='{} m'.format(_fnum(perimetro,',.2f'))
    _set_cell_bg(dr.cells[0],'EAF4EA'); _set_cell_bg(dr.cells[1],'EAF4EA')
    for c in dr.cells:
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs:
            r.font.name='Arial'; r.font.size=Pt(12); r.font.bold=True
        _set_cell_borders(c,'2E6E3E',4)
    doc.add_paragraph().paragraph_format.space_after=Pt(8)


def _seccion_mapa(doc, info):
    _heading(doc, 'VI.  INFORMACIÓN TÉCNICA DEL MAPA', 1)
    _para_just(doc,
        'El presente plano ha sido elaborado utilizando el sistema de referencia '
        'geodésico y la proyección cartográfica indicados a continuación:',
        sp_before=4, sp_after=6)
    if info:
        t=doc.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
        hdr=t.rows[0]; hdr.cells[0].text='PARÁMETRO'; hdr.cells[1].text='VALOR'
        _fmt_hdr(hdr)
        for i,(d,v) in enumerate(info.items()):
            if v:
                row=t.add_row(); row.cells[0].text=str(d); row.cells[1].text=str(v)
                _fmt_row(row, i%2==1)
                for c in row.cells:
                    c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
                    for r in c.paragraphs[0].runs: r.font.name='Arial'; r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(8)


def _seccion_firma(doc, dat):
    p_sep=doc.add_paragraph(); p_sep.paragraph_format.space_before=Pt(18)
    _add_bottom_border(p_sep,'17375E','4')

    p_f=doc.add_paragraph(); p_f.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p_f.paragraph_format.space_before=Pt(8); p_f.paragraph_format.space_after=Pt(28)
    r=p_f.add_run('Puerto Maldonado, {}'.format(_fecha_es()))
    r.font.name='Arial'; r.font.size=Pt(10); r.font.italic=True

    for _ in range(4):
        sp=doc.add_paragraph(''); sp.paragraph_format.line_spacing=Pt(13)

    for txt, bold, size, color in [
        ('_'*40, False, 11, None),
        ((dat.get('nombre','') or 'SOLICITANTE').upper(), True, 11, None),
        ('D.N.I. N\u00b0 {}'.format(dat.get('dni','')), False, 10, None),
        ('PROPIETARIO / SOLICITANTE', True, 10, RGBColor(0x17,0x37,0x5E)),
    ]:
        p=doc.add_paragraph(txt); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2)
        for r in p.runs:
            r.font.name='Arial'; r.font.size=Pt(size); r.font.bold=bold
            if color: r.font.color.rgb=color
